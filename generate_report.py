import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

def generate_full_data_report(base_root, case_letter):
    """
    Case A or B folders are scanned to collect all images, Excel/CSV (full rows), and text files.
    """
    main_folders = [d for d in os.listdir(base_root) 
                    if os.path.isdir(os.path.join(base_root, d)) and f"case{case_letter}" in d]
    main_folders.sort()

    if not main_folders:
        print(f"Warning: No folders found for Case {case_letter}.")
        return

    doc = Document()
    doc.add_heading(f'Motor Fault Diagnosis Report - CASE {case_letter} (Full Data)', 0)

    for main_f in main_folders:
        main_folder_path = os.path.join(base_root, main_f)
        sub_folders = sorted([d for d in os.listdir(main_folder_path) 
                             if os.path.isdir(os.path.join(main_folder_path, d))])
        
        for sub_f in sub_folders:
            folder_path = os.path.join(main_folder_path, sub_f)
            
            # --- Section Title ---
            set_label = f"SECTION: {main_f.upper()} / {sub_f}"
            p = doc.add_paragraph()
            run = p.add_run(f"■ {set_label}")
            run.bold = True
            run.font.size = Pt(14)

            files = os.listdir(folder_path)
            
            # 1. Image files
            images = sorted([f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg'))])
            for img in images:
                img_path = os.path.join(folder_path, img)
                doc.add_paragraph(f"  [Image] {img}").runs[0].italic = True
                doc.add_picture(img_path, width=Inches(4.2))

            # 2. Excel / CSV files (all rows included)
            data_files = sorted([f for f in files if f.lower().endswith(('.csv', '.xlsx', '.xls'))])
            
            for d_file in data_files:
                file_path = os.path.join(folder_path, d_file)
                doc.add_paragraph(f"  [Full Data Table] {d_file}").runs[0].bold = True
                
                try:
                    if d_file.lower().endswith('.csv'):
                        df = pd.read_csv(file_path)
                    else:
                        df = pd.read_excel(file_path)
                    
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = 'Table Grid'
                    
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(df.columns):
                        hdr_cells[i].text = str(col)
                    
                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)
                            
                except Exception as e:
                    doc.add_paragraph(f"  (Error reading file: {d_file} - {e})")
                
                doc.add_paragraph()

            # 3. Text files
            txt_files = sorted([f for f in files if f.lower().endswith('.txt')])
            for txt in txt_files:
                doc.add_paragraph(f"  [Text Info] {txt}").runs[0].italic = True
                txt_path = os.path.join(folder_path, txt)
                with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                    doc.add_paragraph(f.read())

            doc.add_page_break()

    output_name = f"Full_Report_Case{case_letter}.docx"
    doc.save(os.path.join(base_root, output_name))
    print(f"Done: {output_name} has been created.")

if __name__ == "__main__":
    # Set your project folder path here
    BASE_PATH = r"your\project\folder\path"
    
    generate_full_data_report(BASE_PATH, "A")
    generate_full_data_report(BASE_PATH, "B")
